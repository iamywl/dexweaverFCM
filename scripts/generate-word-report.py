#!/usr/bin/env python3
"""
모든 프로젝트 문서를 하나의 Word 파일로 통합 생성
실행: python3 scripts/generate-word-report.py
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ─── 문서 구조 정의 ───
PARTS = [
    {
        "title": "제1부  연구 기반",
        "chapters": [
            ("제1장  FCM 기술 조사", "docs/01-fcm-research.md"),
            ("제2장  FCM 케이스 분석", "docs/02-fcm-case-analysis.md"),
            ("제3장  QoS 평가 지표 정의", "docs/03-qos-evaluation-metrics.md"),
        ],
    },
    {
        "title": "제2부  실험 설계 및 개발 계획",
        "chapters": [
            ("제4장  실험 설계", "docs/04-experiment-design.md"),
            ("제5장  개발 계획", "docs/05-development-plan.md"),
            ("제6장  서버리스 구현 계획", "docs/06-serverless-implementation-plan.md"),
            ("제7장  서버리스 개발 계획", "docs/07-serverless-dev-plan.md"),
        ],
    },
    {
        "title": "제3부  실험 결과",
        "chapters": [
            ("제8장  서버→폰 QoS 분석 보고서 (그룹 A~F)", "results/server-to-phone-report.md"),
            ("제9장  추가 실험 QoS 분석 보고서 (그룹 G~J)", "results/additional-report.md"),
            ("제10장  Phase 1 QoS 메트릭 보고서", "results/qos-metrics-report.md"),
        ],
    },
    {
        "title": "제4부  평가 체계 및 운영 가이드",
        "chapters": [
            ("제11장  QoS 평가 지표 상세", "evaluation/01-evaluation-metrics.md"),
            ("제12장  FCM 시나리오 분류", "evaluation/02-fcm-scenarios.md"),
            ("제13장  FCM 알림 신뢰성 확보 방안", "evaluation/03-reliability-guide.md"),
        ],
    },
    {
        "title": "부록",
        "chapters": [
            ("부록 A  환경 구축 진행 보고서", "reports/01-setup-progress-report.md"),
            ("부록 B  프로젝트 진행 보고서", "docs/progress-report.md"),
        ],
    },
]


def read_md(relpath):
    fpath = os.path.join(BASE, relpath)
    if not os.path.exists(fpath):
        return f"(파일 없음: {relpath})"
    with open(fpath, "r", encoding="utf-8") as f:
        return f.read()


# ─── 스타일 설정 ───

def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Heading 스타일
    for level in range(1, 5):
        sname = f"Heading {level}"
        if sname in doc.styles:
            s = doc.styles[sname]
            s.font.name = "맑은 고딕"
            s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            if level == 1:
                s.font.size = Pt(18)
                s.font.bold = True
            elif level == 2:
                s.font.size = Pt(14)
                s.font.bold = True
            elif level == 3:
                s.font.size = Pt(12)
                s.font.bold = True
            elif level == 4:
                s.font.size = Pt(11)
                s.font.bold = True


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DexWeaver FCM")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("QoS 연구 보고서")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x66)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Firebase Cloud Messaging 구현 방법에 따른\n서비스 품질(QoS) 정량 평가 및 신뢰성 확보 방안")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026년 3월")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("환경: Android 에뮬레이터 (Pixel 3a API 36), Firebase Spark 플랜")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("목차", level=1)
    doc.add_paragraph("")

    for part in PARTS:
        p = doc.add_paragraph()
        run = p.add_run(part["title"])
        run.font.bold = True
        run.font.size = Pt(11)

        for ch_title, _ in part["chapters"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(ch_title)
            run.font.size = Pt(10)

        doc.add_paragraph("")

    doc.add_page_break()


# ─── Markdown → Word 변환 ───

def parse_table(lines):
    """마크다운 테이블을 2D 리스트로 파싱"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # 구분선(---|---) 스킵
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if all(re.match(r"^[\-:]+$", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows or len(rows) < 1:
        return
    ncols = max(len(r) for r in rows)

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            # 볼드 처리 제거하고 텍스트만
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            clean = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", clean)  # 링크 제거
            run = p.add_run(clean)
            run.font.size = Pt(8)
            run.font.name = "맑은 고딕"
            if i == 0:
                run.font.bold = True
                # 헤더 배경색
                shading = cell._element.get_or_add_tcPr()
                bg = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): "E8EAF6",
                    qn("w:val"): "clear",
                })
                shading.append(bg)

    doc.add_paragraph("")


def add_md_content(doc, md_text, skip_first_heading=True):
    """마크다운 텍스트를 Word에 추가"""
    lines = md_text.split("\n")
    i = 0
    first_heading_skipped = False
    table_buffer = []
    in_table = False
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # 코드 블록
        if line.strip().startswith("```"):
            if in_code_block:
                # 코드 블록 종료
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_text)
                run.font.size = Pt(8)
                run.font.name = "Courier New"
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 테이블 처리
        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        elif table_buffer:
            rows = parse_table(table_buffer)
            if rows:
                add_table(doc, rows)
            table_buffer = []

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 구분선
        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        # 각주
        if line.strip().startswith("[^"):
            p = doc.add_paragraph()
            clean = re.sub(r"\[(\^[0-9]+)\]:\s*", r"[\1] ", line.strip())
            run = p.add_run(clean)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 헤딩
        heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            if skip_first_heading and not first_heading_skipped and level == 1:
                first_heading_skipped = True
                i += 1
                continue

            # 마크다운 서식 제거
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)

            doc.add_heading(text, level=min(level + 1, 4))
            i += 1
            continue

        # 인용구
        if line.strip().startswith(">"):
            text = re.sub(r"^>\s*", "", line.strip())
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            run.font.size = Pt(9)
            i += 1
            continue

        # 리스트
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)", line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            text = re.sub(r"\*(.*?)\*", r"\1", text)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)

            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1 + indent * 0.5)
            run = p.add_run(text)
            run.font.size = Pt(9)
            i += 1
            continue

        # 일반 텍스트
        text = line.strip()
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)

        if text:
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(10) if p.runs else None
        i += 1

    # 남은 테이블 처리
    if table_buffer:
        rows = parse_table(table_buffer)
        if rows:
            add_table(doc, rows)


# ─── 메인 ───

def main():
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    setup_styles(doc)
    add_title_page(doc)
    add_toc(doc)

    for part in PARTS:
        # 파트 제목 페이지
        for _ in range(4):
            doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(part["title"])
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        doc.add_page_break()

        for ch_title, relpath in part["chapters"]:
            doc.add_heading(ch_title, level=1)
            p = doc.add_paragraph()
            run = p.add_run(f"원본: {relpath}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            doc.add_paragraph("")

            md = read_md(relpath)
            add_md_content(doc, md, skip_first_heading=True)
            doc.add_page_break()

    # 저장
    out_path = os.path.join(BASE, "DexWeaver_FCM_QoS_연구보고서.docx")
    doc.save(out_path)
    print(f"✅ Word 문서 생성 완료: {out_path}")
    print(f"   파일 크기: {os.path.getsize(out_path) / 1024:.0f} KB")


if __name__ == "__main__":
    main()
